"""Insert the approved three-tier control figures into a copy of the paper."""

from __future__ import annotations

import hashlib
import os
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


SOURCE = Path(
    r"C:\Users\31541\Desktop\油水井\高采三区数智化注采管理系统的构建与应用(高升采油厂采三)7.13-完善版4.docx"
)
ARTIFACT_DIR = Path(__file__).resolve().parent
OUTPUT = ARTIFACT_DIR / "高采三区数智化注采管理系统的构建与应用-三级管控插图完成版.docx"
TRUSTED_SOURCE_SHA256 = "c5b356318a4484bc18e3e54325df1321332ae7f3c01a0497fb5ec1347479ba32"
FIGURES = {
    "[系统三级管控架构图]": ARTIFACT_DIR / "figure-1-three-tier-control-architecture.png",
    "[三级协同管控示例图]": ARTIFACT_DIR / "figure-2-three-tier-collaboration-example.png",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> None:
    source_hash_before = sha256(SOURCE)
    if source_hash_before != TRUSTED_SOURCE_SHA256:
        raise RuntimeError(
            "Source DOCX does not match the trusted SHA-256: "
            f"expected {TRUSTED_SOURCE_SHA256}, got {source_hash_before}"
        )

    if SOURCE.resolve() == OUTPUT.resolve():
        raise RuntimeError("Output path must differ from the source DOCX")
    try:
        if OUTPUT.exists() and os.path.samefile(SOURCE, OUTPUT):
            raise RuntimeError("Output DOCX must not be a hard link alias of the source DOCX")
    except FileNotFoundError:
        pass

    document = Document(SOURCE)
    paragraphs = document.paragraphs
    occurrences = {placeholder: 0 for placeholder in FIGURES}

    for paragraph in paragraphs:
        placeholder = paragraph.text.strip()
        if placeholder in occurrences:
            occurrences[placeholder] += 1

    invalid = {placeholder: count for placeholder, count in occurrences.items() if count != 1}
    if invalid:
        details = ", ".join(
            f"{placeholder!r}: expected 1, found {count}" for placeholder, count in invalid.items()
        )
        raise RuntimeError(f"Placeholder occurrence mismatch ({details})")

    for paragraph in paragraphs:
        placeholder = paragraph.text.strip()
        if placeholder not in FIGURES:
            continue
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.add_run().add_picture(str(FIGURES[placeholder]), width=Mm(146))

    temporary_output: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            prefix=f".{OUTPUT.stem}-",
            suffix=".docx",
            dir=OUTPUT.parent,
            delete=False,
        ) as temporary_file:
            temporary_output = Path(temporary_file.name)

        document.save(temporary_output)
        source_hash_after = sha256(SOURCE)
        if source_hash_after != source_hash_before or source_hash_after != TRUSTED_SOURCE_SHA256:
            raise RuntimeError(
                "Source DOCX changed during figure insertion: "
                f"before {source_hash_before}, after {source_hash_after}"
            )
        os.replace(temporary_output, OUTPUT)
    finally:
        if temporary_output is not None:
            try:
                temporary_output.unlink()
            except FileNotFoundError:
                pass


if __name__ == "__main__":
    main()
