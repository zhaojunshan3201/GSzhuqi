"""Acceptance tests for the three-tier control paper artifacts."""

from __future__ import annotations

import hashlib
import os
import unittest
import xml.etree.ElementTree as ET
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from PIL import Image, ImageChops


ARTIFACT_DIR = Path(__file__).resolve().parent
FIGURE_1_SVG = ARTIFACT_DIR / "figure-1-three-tier-control-architecture.svg"
FIGURE_2_SVG = ARTIFACT_DIR / "figure-2-three-tier-collaboration-example.svg"
FIGURE_1_PNG = ARTIFACT_DIR / "figure-1-three-tier-control-architecture.png"
FIGURE_2_PNG = ARTIFACT_DIR / "figure-2-three-tier-collaboration-example.png"
OUTPUT_DOCX = ARTIFACT_DIR / "高采三区数智化注采管理系统的构建与应用-三级管控插图完成版.docx"
DEFAULT_SOURCE_DOCX = Path(
    r"C:\Users\31541\Desktop\油水井\高采三区数智化注采管理系统的构建与应用(高升采油厂采三)7.13-完善版4.docx"
)
# PAPER_SOURCE_DOCX may relocate an identical, trusted source DOCX for another machine.
TRUSTED_SOURCE_SHA256 = "c5b356318a4484bc18e3e54325df1321332ae7f3c01a0497fb5ec1347479ba32"


def source_docx_path() -> Path:
    """Use PAPER_SOURCE_DOCX when set; otherwise use the project source path."""
    return Path(os.environ.get("PAPER_SOURCE_DOCX", str(DEFAULT_SOURCE_DOCX)))


class ThreeTierArtifactTests(unittest.TestCase):
    def require_file(self, path: Path) -> None:
        if not path.is_file():
            self.fail(f"Required artifact is missing: {path}")

    def assert_svg(self, path: Path, view_box: str, required_text: list[str]) -> None:
        self.require_file(path)
        root = ET.parse(path).getroot()
        self.assertEqual(root.attrib.get("viewBox"), view_box)

        text_nodes = [node for node in root.iter() if node.tag.rsplit("}", 1)[-1] == "text"]
        self.assertTrue(text_nodes, f"SVG has no text nodes: {path}")
        for node in text_nodes:
            self.assertTrue(
                "".join(node.itertext()).strip(),
                f"SVG contains an empty text node: {path}",
            )

        visible_text = "".join("".join(node.itertext()) for node in text_nodes)
        for phrase in required_text:
            self.assertIn(phrase, visible_text, f"Missing visible SVG text {phrase!r} in {path}")

    def assert_png(self, path: Path, expected_height: int) -> None:
        self.require_file(path)
        with Image.open(path) as image:
            self.assertIn(image.mode, {"RGB", "RGBA"})
            self.assertEqual(image.size, (2400, expected_height))
            for point in ((0, 0), (2399, 0), (0, expected_height - 1), (2399, expected_height - 1)):
                pixel = image.getpixel(point)
                self.assertEqual(pixel[:3], (255, 255, 255), f"Corner {point} is not white in {path}")
            white_background = Image.new("RGB", image.size, "white")
            self.assertIsNotNone(
                ImageChops.difference(image.convert("RGB"), white_background).getbbox(),
                f"PNG is entirely white: {path}",
            )

    def assert_caption_follows_png(self, document: Document, caption: str, png_path: Path) -> None:
        self.require_file(png_path)
        paragraphs = document.paragraphs
        caption_indexes = [index for index, paragraph in enumerate(paragraphs) if paragraph.text == caption]
        self.assertEqual(len(caption_indexes), 1, f"Missing or duplicate caption: {caption}")
        self.assertGreater(caption_indexes[0], 0, f"Caption has no preceding paragraph: {caption}")

        image_paragraph = paragraphs[caption_indexes[0] - 1]
        drawing = image_paragraph._p.find(".//" + qn("w:drawing"))
        self.assertIsNotNone(drawing, f"Caption does not immediately follow an image paragraph: {caption}")
        blip = drawing.find(".//" + qn("a:blip"))
        self.assertIsNotNone(blip, f"Image paragraph has no embedded image: {caption}")
        relationship_id = blip.get(qn("r:embed"))
        self.assertTrue(relationship_id, f"Embedded image has no relationship id: {caption}")
        embedded_image = document.part.rels[relationship_id].target_part.blob
        embedded_hash = hashlib.sha256(embedded_image).hexdigest()
        expected_hash = hashlib.sha256(png_path.read_bytes()).hexdigest()
        self.assertEqual(embedded_hash, expected_hash, f"Caption is paired with the wrong image: {caption}")

    def test_figure_1_svg(self) -> None:
        self.assert_svg(
            FIGURE_1_SVG,
            "0 0 1120 620",
            [
                "注汽全流程数字化闭环", "区块级", "井组级", "单井级", "诊断信息逐级上推",
                "调控决策逐级反馈", "统一数据底座", "Oracle / Excel", "SQLite缓存", "井号（JH）",
            ],
        )

    def test_figure_2_svg(self) -> None:
        self.assert_svg(
            FIGURE_2_SVG,
            "0 0 1120 590",
            [
                "单井精细诊断", "井组注采平衡", "区块整体调控", "+6 个百分点", "高3624块典型井组",
                "17 井次", "−195 t", "−215 t", "+15%", "高246区块拉大井距", "2200 → 2050 t",
                "选井效率提升 80% 以上", "异常响应 15–30 天 → 1–3 天",
            ],
        )

    def test_figure_1_png(self) -> None:
        self.assert_png(FIGURE_1_PNG, 1329)

    def test_figure_2_png(self) -> None:
        self.assert_png(FIGURE_2_PNG, 1264)

    def test_completed_docx_replaces_placeholders_and_adds_figures(self) -> None:
        source_docx = source_docx_path()
        self.require_file(source_docx)
        self.require_file(OUTPUT_DOCX)
        self.assertNotEqual(source_docx.resolve(), OUTPUT_DOCX.resolve())
        self.assertFalse(os.path.samefile(source_docx, OUTPUT_DOCX), "Output DOCX must not be the source DOCX or its hard link")

        source_hash_before = hashlib.sha256(source_docx.read_bytes()).hexdigest()
        self.assertEqual(source_hash_before, TRUSTED_SOURCE_SHA256, "Source DOCX does not match the trusted baseline")
        source = Document(source_docx)
        output = Document(OUTPUT_DOCX)
        output_text = "\n".join(paragraph.text for paragraph in output.paragraphs)

        self.assertNotIn("[系统三级管控架构图]", output_text)
        self.assertNotIn("[三级协同管控示例图]", output_text)
        self.assertIn("图1  系统三级管控架构示意图", output_text)
        self.assertIn("图2  井→组→区块三级协同管控效果示例", output_text)
        self.assertGreaterEqual(len(output.inline_shapes), len(source.inline_shapes) + 2)
        self.assert_caption_follows_png(output, "图1  系统三级管控架构示意图", FIGURE_1_PNG)
        self.assert_caption_follows_png(output, "图2  井→组→区块三级协同管控效果示例", FIGURE_2_PNG)

        source_hash_after = hashlib.sha256(source_docx.read_bytes()).hexdigest()
        self.assertEqual(source_hash_before, source_hash_after, "Source DOCX changed during verification")


if __name__ == "__main__":
    unittest.main(verbosity=2)
